"""Generate a recommendation letter as .docx matching George's style."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def set_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def make_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   space_after=6, space_before=0, font_size=12,
                   bold=False, italic=False, first_line_indent=None):
    """Add a paragraph with consistent formatting."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.0

    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)

    run = p.add_run(text)
    set_font(run, size=font_size, bold=bold, italic=italic)
    return p


def generate_letter(
    output_path,
    student_name,
    date_str,
    recipient,
    subject,
    body_paragraphs,
    signature_name,
    signature_title,
    signature_school,
    signature_email=None,
    signature_phone=None,
):
    """
    Generate a .docx recommendation letter.

    Args:
        output_path: Full path for the .docx file
        student_name: Student's name (for filename, not used in content)
        date_str: e.g. "May 19, 2026"
        recipient: e.g. "To the Admissions Committee,"
        subject: Optional subject line (can be empty string)
        body_paragraphs: List of paragraph strings (no newlines within)
        signature_name: "Yuqi (George) Gu"
        signature_title: e.g. "Associate Professor of Finance"
        signature_school: e.g. "Atkinson Graduate School of Management\nWillamette University"
        signature_email: e.g. "ygu@willamette.edu"
        signature_phone: e.g. "503-370-6883"
    """
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # --- Date (right-aligned) ---
    make_paragraph(doc, date_str, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                   space_after=12)

    # --- Recipient ---
    make_paragraph(doc, recipient, space_after=6)

    # --- Subject (optional) ---
    if subject:
        make_paragraph(doc, subject, bold=True, space_after=12)

    # --- Salutation ---
    def extract_committee(recipient):
        """Extract salutation from recipient."""
        if "Admissions Committee" in recipient or "admissions committee" in recipient:
            return "To the Admissions Committee,"
        if "Scholarship Committee" in recipient:
            return "To the Scholarship Committee,"
        return "To Whom It May Concern:"

    make_paragraph(doc, extract_committee(recipient), space_after=12)

    # --- Body paragraphs ---
    for para_text in body_paragraphs:
        make_paragraph(doc, para_text, space_after=12)

    # --- Closing ---
    make_paragraph(doc, "Sincerely,", space_before=12, space_after=24)

    # --- Signature (3 blank lines' worth of space) ---
    for _ in range(3):
        make_paragraph(doc, "", space_after=0)

    # --- Name ---
    make_paragraph(doc, signature_name, bold=True, space_after=2)

    # --- Title & School ---
    make_paragraph(doc, signature_title, space_after=0)
    for line in signature_school.split("\n"):
        make_paragraph(doc, line, space_after=0)

    # --- Contact ---
    contact_parts = []
    if signature_email:
        contact_parts.append(signature_email)
    if signature_phone:
        contact_parts.append(signature_phone)
    if contact_parts:
        make_paragraph(doc, " | ".join(contact_parts), space_after=0)

    # --- Save ---
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
